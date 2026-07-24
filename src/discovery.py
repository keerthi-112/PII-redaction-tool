# Placeholder: use finalized discovery implementation from our discussion.
# Placeholder: use finalized discovery implementation from our discussion.

from __future__ import annotations
import re
from typing import List

import spacy
from docx import Document
from presidio_analyzer import AnalyzerEngine

from config import SPACY_MODEL

from detector_patterns import (
    EMAIL_REGEX,
    PHONE_REGEX,
    SSN_REGEX,
    CREDIT_CARD_REGEX,
    IP_REGEX,
    DOB_REGEX,
)
ORG_REGEX = re.compile(
    r"\b[A-Z][A-Za-z0-9&.,'-]*(?:\s+[A-Z][A-Za-z0-9&.,'-]*)*\s+"
    r"(?:Corporation|Corp|Inc|Incorporated|Ltd|Limited|LLC|PLC|Company|Co)\b"
)

from utils import DetectedEntity


class Discovery:

    def __init__(self):

        print("Loading spaCy model...")
        self.nlp = spacy.load(SPACY_MODEL)

        print("Loading Presidio...")
        self.analyzer = AnalyzerEngine()

    # =====================================================
    # Regex Detection
    # =====================================================

    def _regex_detection(
        self,
        text,
        **location
    ):

        entities = []

        patterns = [
            
            ("EMAIL", EMAIL_REGEX),
            
            ("SSN", SSN_REGEX),
            
            ("CREDIT_CARD", CREDIT_CARD_REGEX),
            
            ("IP_ADDRESS", IP_REGEX),
            
            ("DATE_OF_BIRTH", DOB_REGEX),
            
            ("PHONE_NUMBER", PHONE_REGEX),
        
        ]

        for label, pattern in patterns:

            for match in pattern.finditer(text):

                entities.append(

                    DetectedEntity(

                        text=match.group(),

                        label=label,

                        start=match.start(),

                        end=match.end(),

                        confidence=1.0,

                        source="Regex",

                        **location

                    )

                )

        return entities

    # =====================================================
    # spaCy Detection
    # =====================================================

    def _spacy_detection(
        self,
        text,
        **location
    ):

        entities = []

        doc = self.nlp(text)

        label_map = {
            "PERSON": "PERSON",
            "ORG": "ORGANIZATION",
            "GPE": "ADDRESS",
            "LOC": "ADDRESS",
        }

        IGNORE_WORDS = {
            "name",
            "email",
            "phone",
            "ssn",
            "address",
            "organization",
            "credit",
            "card",
            "credit card",
            "ip",
            "ip address",
            "date of birth",
        }
    # -----------------------------
    # spaCy entities
    # -----------------------------
        for ent in doc.ents:

            if ent.label_ not in label_map:
                continue

            value = ent.text.strip()

            if not value:
                continue

            if value.lower() in IGNORE_WORDS:
                continue

            entities.append(
                DetectedEntity(
                    text=value,
                    label=label_map[ent.label_],
                    start=ent.start_char,
                    end=ent.end_char,
                    confidence=0.95,
                    source="spaCy",
                    **location
                )
            )

    # -----------------------------
    # Organization regex fallback
    # -----------------------------

        for match in ORG_REGEX.finditer(text):

            value = match.group().strip()

            if any(
                e.start == match.start()
                and e.end == match.end()
                for e in entities
            ):
                continue

            entities.append(
                DetectedEntity(
                    text=value,
                    label="ORGANIZATION",
                    start=match.start(),
                    end=match.end(),
                    confidence=0.90,
                    source="RegexORG",
                    **location
                )
            )

        return entities

    # =====================================================
    # Presidio Detection
    # =====================================================

    def _presidio_detection(
        self,
        text,
        **location
    ):

        entities = []

        results = self.analyzer.analyze(
            text=text,
            language="en"
        )

        label_map = {

            "PERSON": "PERSON",

            "EMAIL_ADDRESS": "EMAIL",

            "PHONE_NUMBER": "PHONE_NUMBER",

            "ORGANIZATION": "ORGANIZATION",

            "LOCATION": "ADDRESS",

            "IP_ADDRESS": "IP_ADDRESS",

            "US_SSN": "SSN",

            "CREDIT_CARD": "CREDIT_CARD",

            "DATE_TIME": "DATE_OF_BIRTH"

        }
        IGNORE_WORDS = {
                "name",
                "email",
                "phone",
                "ssn",
                "address",
                "organization",
                "credit",
                "card",
                "credit card",
                "ip",
                "ip address",
                "date of birth",
        }

        for result in results:

            if result.entity_type not in label_map:
                continue

            value = text[result.start:result.end].strip()

            if not value:
                continue
           

            
            if value.lower() in IGNORE_WORDS:
                continue

            entities.append(

                DetectedEntity(

                    text=value,

                    label=label_map[result.entity_type],

                    start=result.start,

                    end=result.end,

                    confidence=result.score,

                    source="Presidio",

                    **location

                )

            )

        return entities

    # =====================================================
    # Remove Duplicates
    # =====================================================

    @staticmethod
    def _remove_duplicates(
        entities: List[DetectedEntity]
    ) -> List[DetectedEntity]:

        unique = {}

        for entity in entities:

            key = (

                entity.text,

                entity.label,

                entity.start,

                entity.end,

                entity.paragraph_index,

                entity.table_index,

                entity.row_index,

                entity.cell_index

            )

            if key not in unique:

                unique[key] = entity

                continue

            if entity.confidence > unique[key].confidence:

                unique[key] = entity

        return list(unique.values())

    # =====================================================
    # Resolve Overlapping Entities
    # =====================================================

    @staticmethod
    def _resolve_overlaps(
        entities: List[DetectedEntity]
    ) -> List[DetectedEntity]:

        if not entities:
            return []

        entities.sort(

            key=lambda e: (

                e.start,

                -(e.end - e.start),

                -e.confidence

            )

        )

        resolved = []

        for entity in entities:

            should_add = True

            for existing in resolved:

                same_location = (

                    entity.paragraph_index == existing.paragraph_index

                    and entity.table_index == existing.table_index

                    and entity.row_index == existing.row_index

                    and entity.cell_index == existing.cell_index

                )

                if not same_location:
                    continue

                overlap = (

                    entity.start < existing.end

                    and entity.end > existing.start

                )

                if not overlap:
                    continue


                # If one entity is completely inside another,
                # keep the larger entity.
                if (
                    entity.start >= existing.start
                    and entity.end <= existing.end
                ):
                    should_add = False
                    break

                if (
                    existing.start >= entity.start
                    and existing.end <= entity.end
                ):
                    resolved.remove(existing)
                    resolved.append(entity)
                    should_add = False
                    break

                current_length = entity.end - entity.start

                existing_length = existing.end - existing.start

                if current_length > existing_length:

                    resolved.remove(existing)

                    resolved.append(entity)

                elif (

                    current_length == existing_length

                    and entity.confidence > existing.confidence

                ):

                    resolved.remove(existing)

                    resolved.append(entity)

                should_add = False

                break

            if should_add:

                resolved.append(entity)

        resolved.sort(

            key=lambda e: (

                e.paragraph_index if e.paragraph_index is not None else 999999,

                e.table_index if e.table_index is not None else 999999,

                e.start

            )

        )

        return resolved

        # =====================================================
    # Process Paragraph
    # =====================================================

    def _process_paragraph(
        self,
        paragraph,
        paragraph_index
    ):

        text = paragraph.text.strip()

        if not text:
            return []

        entities = []

        entities.extend(

            self._regex_detection(
                text,
                paragraph_index=paragraph_index
            )

        )

        entities.extend(

            self._spacy_detection(
                text,
                paragraph_index=paragraph_index
            )

        )

        entities.extend(

            self._presidio_detection(
                text,
                paragraph_index=paragraph_index
            )

        )

        return entities

    # =====================================================
    # Process Tables
    # =====================================================

    def _process_tables(
        self,
        document
    ):

        entities = []

        for table_index, table in enumerate(document.tables):

            for row_index, row in enumerate(table.rows):

                for cell_index, cell in enumerate(row.cells):

                    text = cell.text.strip()

                    if not text:
                        continue

                    entities.extend(

                        self._regex_detection(

                            text,

                            table_index=table_index,

                            row_index=row_index,

                            cell_index=cell_index

                        )

                    )

                    entities.extend(

                        self._spacy_detection(

                            text,

                            table_index=table_index,

                            row_index=row_index,

                            cell_index=cell_index

                        )

                    )

                    entities.extend(

                        self._presidio_detection(

                            text,

                            table_index=table_index,

                            row_index=row_index,

                            cell_index=cell_index

                        )

                    )

        return entities

    # =====================================================
    # Discover Entities
    # =====================================================

    def discover(self, input_file):
        print("\nStarting Discovery...\n")

        document = Document(input_file)

        entities = []

        for paragraph_index, paragraph in enumerate(document.paragraphs):

            entities.extend(

                self._process_paragraph(

                    paragraph,

                    paragraph_index

                )

            )

        entities.extend(

            self._process_tables(document)

        )

        print(f"Detected {len(entities)} entities before cleaning.")

        entities = self._remove_duplicates(entities)

        print(f"{len(entities)} entities after duplicate removal.")

        entities = self._resolve_overlaps(entities)

        print(f"{len(entities)} entities after overlap resolution.")

        return document, entities

        # =====================================================
    # Public API
    # =====================================================

    def run(self, input_file):
        """
        Main entry point for the Discovery stage.

        Returns
        -------
        tuple
            (document, detected_entities)
        """

        try:

            document, entities = self.discover(input_file)

            print("\n========== Discovery Summary ==========\n")

            print(f"Total Entities Detected : {len(entities)}")

            counts = {}

            for entity in entities:

                counts[entity.label] = counts.get(entity.label, 0) + 1

            for label, count in sorted(counts.items()):

                print(f"{label:<20} : {count}")

            print("\n=======================================\n")

            return document, entities

        except Exception as e:

            raise RuntimeError(

                f"Discovery stage failed: {str(e)}"

            ) from e