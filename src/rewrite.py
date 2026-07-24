from __future__ import annotations

from typing import List

from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .utils import DetectedEntity


class Rewrite:
    """
    Rewrites detected entities inside a DOCX while preserving
    formatting as much as possible.

    The implementation operates at the run level instead of using
    paragraph.text replacement, which would destroy formatting.
    """

    def __init__(self):

        pass

    # =====================================================
    # Helpers
    # =====================================================

    @staticmethod
    def _paragraph_text(paragraph: Paragraph) -> str:

        return "".join(run.text for run in paragraph.runs)

    @staticmethod
    def _cell_text(cell: _Cell) -> str:

        return "\n".join(

            "".join(run.text for run in paragraph.runs)

            for paragraph in cell.paragraphs

        )

    # =====================================================
    # Replace inside a paragraph
    # =====================================================

    def _replace_in_paragraph(

        self,

        paragraph: Paragraph,

        entities: List[DetectedEntity]

    ) -> None:

        if not paragraph.runs:

            return

        text = self._paragraph_text(paragraph)

        if not text:

            return

        entities = sorted(

            entities,

            key=lambda e: e.start,

            reverse=True

        )

        for entity in entities:

            original = entity.text

            replacement = entity.replacement

            start = entity.start

            end = entity.end

            if start >= len(text):

                continue

            if text[start:end] != original:

                continue

            current = 0

            start_run = None

            end_run = None

            start_offset = 0

            end_offset = 0

            for index, run in enumerate(paragraph.runs):

                run_start = current

                run_end = current + len(run.text)

                if start_run is None and run_start <= start < run_end:

                    start_run = index

                    start_offset = start - run_start

                if run_start <= end <= run_end:

                    end_run = index

                    end_offset = end - run_start

                    break

                current = run_end

            if start_run is None or end_run is None:

                continue

            if start_run == end_run:

                run = paragraph.runs[start_run]

                run.text = (

                    run.text[:start_offset]

                    + replacement

                    + run.text[end_offset:]

                )

                continue

            first = paragraph.runs[start_run]

            last = paragraph.runs[end_run]

            prefix = first.text[:start_offset]

            suffix = last.text[end_offset:]

            first.text = prefix + replacement

            for i in range(start_run + 1, end_run):

                paragraph.runs[i].text = ""

            last.text = suffix

                # =====================================================
    # Rewrite Paragraphs
    # =====================================================

    def _rewrite_paragraphs(
        self,
        document: Document,
        entities: List[DetectedEntity]
    ) -> None:

        paragraph_map = {}

        for entity in entities:

            if entity.paragraph_index is None:
                continue

            paragraph_map.setdefault(
                entity.paragraph_index,
                []
            ).append(entity)

        for index, paragraph in enumerate(document.paragraphs):

            if index not in paragraph_map:
                continue

            self._replace_in_paragraph(
                paragraph,
                paragraph_map[index]
            )

    # =====================================================
    # Rewrite Tables
    # =====================================================

    def _rewrite_tables(
        self,
        document: Document,
        entities: List[DetectedEntity]
    ) -> None:

        table_map = {}

        for entity in entities:

            if entity.table_index is None:
                continue

            key = (
                entity.table_index,
                entity.row_index,
                entity.cell_index
            )

            table_map.setdefault(
                key,
                []
            ).append(entity)

        for table_index, table in enumerate(document.tables):

            for row_index, row in enumerate(table.rows):

                for cell_index, cell in enumerate(row.cells):

                    key = (
                        table_index,
                        row_index,
                        cell_index
                    )

                    if key not in table_map:
                        continue

                    cell_entities = table_map[key]

                    for paragraph in cell.paragraphs:

                        self._replace_in_paragraph(
                            paragraph,
                            cell_entities
                        )

    # =====================================================
    # Rewrite
    # =====================================================

    def rewrite(
        self,
        document: Document,
        entities: List[DetectedEntity]
    ) -> Document:

        print("\nRewriting document...\n")

        self._rewrite_paragraphs(
            document,
            entities
        )

        self._rewrite_tables(
            document,
            entities
        )

        return document

    # =====================================================
    # Public API
    # =====================================================

    def run(
        self,
        document: Document,
        entities: List[DetectedEntity]
    ) -> Document:

        try:

            document = self.rewrite(
                document,
                entities
            )

            print(
                "Document rewritten successfully."
            )

            return document

        except Exception as exc:

            raise RuntimeError(

                f"Rewrite stage failed: {exc}"

            ) from exc